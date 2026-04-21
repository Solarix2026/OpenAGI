# Copyright (c) 2026 ApeironAILab
# OpenAGI — Autonomous Intelligence System
# MIT License — https://github.com/ApeironAILab/OpenAGI

"""
document_reader.py — Read and analyze documents: Word, Excel, PDF, CSV

Uses python-docx, openpyxl, pdfplumber for extraction.
Then NVIDIA NIM analyzes the content.

Tools registered:
- read_document(path, question) → answer question about document
- summarize_document(path) → summarize document
"""
import logging
from pathlib import Path

log = logging.getLogger("DocReader")

SUPPORTED = {
    ".docx": "word",
    ".xlsx": "excel",
    ".xls": "excel",
    ".csv": "csv",
    ".pdf": "pdf",
    ".txt": "text",
    ".md": "text"
}


class DocumentReader:
    def read_word(self, path: str) -> dict:
        """Extract text from Word document (.docx)."""
        try:
            import docx
            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                tables.append(rows)
            return {
                "success": True,
                "text": "\n\n".join(paragraphs),
                "tables": tables,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "path": path
            }
        except ImportError:
            return {"success": False, "error": "pip install python-docx"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def read_excel(self, path: str) -> dict:
        """Extract data from Excel (.xlsx/.xls/.csv)."""
        suffix = Path(path).suffix.lower()
        try:
            if suffix == ".csv":
                import csv
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    reader = csv.DictReader(f)
                    rows = list(reader)
                return {
                    "success": True,
                    "rows": rows[:200],  # Limit for context
                    "columns": list(rows[0].keys()) if rows else [],
                    "row_count": len(rows),
                    "path": path
                }
            else:
                import openpyxl
                wb = openpyxl.load_workbook(path, data_only=True)
                sheets = {}
                for sheet_name in wb.sheetnames[:3]:  # Max 3 sheets
                    ws = wb[sheet_name]
                    data = []
                    for row in ws.iter_rows(max_row=200, values_only=True):
                        if any(c is not None for c in row):
                            data.append([str(c) if c is not None else "" for c in row])
                    sheets[sheet_name] = data
                return {
                    "success": True,
                    "sheets": sheets,
                    "sheet_names": wb.sheetnames,
                    "path": path
                }
        except ImportError as e:
            return {"success": False, "error": f"pip install openpyxl: {e}"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def read_pdf(self, path: str) -> dict:
        """Extract text from PDF."""
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                pages = []
                for i, page in enumerate(pdf.pages[:20]):  # Max 20 pages
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append({"page": i+1, "text": text.strip()})
                full_text = "\n\n---\n\n".join(p["text"] for p in pages)
            return {
                "success": True,
                "text": full_text[:5000],  # Limit for context
                "page_count": len(pages),
                "path": path
            }
        except ImportError:
            return {"success": False, "error": "pip install pdfplumber"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def read_any(self, path: str) -> dict:
        """Auto-detect file type and read."""
        suffix = Path(path).suffix.lower()
        kind = SUPPORTED.get(suffix)
        if not kind:
            return {"success": False, "error": f"Unsupported: {suffix}. Use: {list(SUPPORTED.keys())}"}
        if kind == "word":
            return self.read_word(path)
        if kind == "excel":
            return self.read_excel(path)
        if kind == "pdf":
            return self.read_pdf(path)
        if kind == "text":
            text = Path(path).read_text(encoding="utf-8", errors="replace")
            return {"success": True, "text": text[:5000], "path": path}
        return {"success": False, "error": "Unknown handler"}

    def analyze_document(self, path: str, question: str) -> dict:
        """Read document and answer question using NVIDIA."""
        from core.llm_gateway import call_nvidia

        content = self.read_any(path)
        if not content.get("success"):
            return content

        text = content.get("text", "")
        if not text:
            sheets = content.get("sheets", {})
            text = "\n".join(
                f"Sheet '{name}':\n" + "\n".join("|".join(row) for row in rows[:20])
                for name, rows in sheets.items()
            )

        prompt = f"""Document: {path}
Content:
{text[:4000]}
Question: {question}
Answer based on the document content. Be specific and reference actual data."""

        answer = call_nvidia([{"role": "user", "content": prompt}], max_tokens=600)
        return {
            "success": True,
            "answer": answer,
            "document": path,
            "question": question
        }

    def summarize_document(self, path: str) -> dict:
        """Summarize any document."""
        return self.analyze_document(path, "Summarize the key points and main content.")

    def register_as_tool(self, registry):
        reader = self

        def read_document(params: dict) -> dict:
            path = params.get("path", "")
            question = params.get("question", "")
            if not path:
                return {"success": False, "error": "File path required"}
            if question:
                return reader.analyze_document(path, question)
            return reader.read_any(path)

        def summarize_doc(params: dict) -> dict:
            path = params.get("path", "")
            if not path:
                return {"success": False, "error": "File path required"}
            return reader.summarize_document(path)

        registry.register(
            "read_document",
            read_document,
            "Read and analyze documents: Word (.docx), Excel (.xlsx), PDF, CSV, text",
            {"path": {"type": "string", "required": True}, "question": {"type": "string", "optional": True}},
            "generation"
        )
        registry.register(
            "summarize_document",
            summarize_doc,
            "Summarize any document",
            {"path": {"type": "string", "required": True}},
            "generation"
        )
